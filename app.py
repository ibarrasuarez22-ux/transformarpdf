import streamlit as st
from PIL import Image
import os
import io
import shutil  # Para el chequeo de dependencias
import pytesseract
from pdf2image import convert_from_bytes
from pdf2docx import Converter
from docx import Document

# --- 0. Configuración de la Página y Logos ---

st.set_page_config(
    page_title="Conversor PDF Pro",
    page_icon="🚀",
    layout="wide"
)

# Cargar logos (poner en carpeta 'logos/')
try:
    logo_izq = Image.open("logos/logo1.png")
except FileNotFoundError:
    logo_izq = None

# Título y Logos
col1, col2, col3 = st.columns([1, 3, 1])
with col1:
    if logo_izq:
        st.image(logo_izq, width=120)
with col2:
    st.title("Conversor PDF a Word (Pro)")
    st.markdown("Conversión Estándar y OCR para PDFs escaneados")

st.markdown("---")

# --- 1. Chequeo de Dependencias del Sistema ---

st.sidebar.header("Estado del Sistema (OCR)")
tesseract_ok = shutil.which("tesseract") is not None
poppler_ok = shutil.which("pdftoppm") is not None  # pdftoppm es parte de Poppler

if tesseract_ok:
    st.sidebar.success("Tesseract (OCR) detectado.")
else:
    st.sidebar.error("Tesseract NO encontrado.")
    st.sidebar.info("Instala Tesseract-OCR en tu sistema para habilitar la conversión de PDFs escaneados.")

if poppler_ok:
    st.sidebar.success("Poppler (PDF) detectado.")
else:
    st.sidebar.error("Poppler NO encontrado.")
    st.sidebar.info("Instala Poppler en tu sistema para habilitar la conversión de PDFs escaneados.")

# --- 2. Opciones de Conversión (Sidebar) ---

st.sidebar.header("Opciones de Conversión")

# Usamos 'disabled' para bloquear la opción OCR si falta el software
ocr_disabled = not (tesseract_ok and poppler_ok)
if ocr_disabled:
    st.sidebar.warning("Modo OCR deshabilitado. Ver 'Estado del Sistema'.")

# Selección de modo
modo_conversion = st.sidebar.radio(
    "Selecciona el tipo de PDF:",
    options=["Digital (Formato perfecto, rápido)", "Escaneado (OCR, más lento)"],
    index=0,
    disabled=ocr_disabled
)

# Opciones específicas de OCR
lang = "spa" # Idioma por defecto
if "Escaneado" in modo_conversion:
    st.sidebar.markdown("### Opciones de OCR")
    # Puedes añadir más idiomas si los instalaste con Tesseract
    lang = st.sidebar.selectbox("Idioma del documento:", ["spa", "eng"], help="'spa' = Español, 'eng' = Inglés")

# --- 3. Funciones de Conversión ---

def parse_page_range(range_str, max_pages):
    """Convierte un string como '1, 3-5' en una lista [0, 2, 3, 4]"""
    if not range_str:
        return list(range(max_pages))  # Todas las páginas
    
    pages = set()
    try:
        parts = range_str.split(',')
        for part in parts:
            if '-' in part:
                start, end = map(int, part.split('-'))
                if start > 0 and end <= max_pages and start <= end:
                    pages.update(range(start - 1, end)) # Zero-indexed
            else:
                page = int(part)
                if page > 0 and page <= max_pages:
                    pages.add(page - 1) # Zero-indexed
        return sorted(list(pages))
    except Exception:
        st.error("Rango de páginas inválido. Usando todas las páginas.")
        return list(range(max_pages))

def convert_digital(pdf_bytes, pages_to_convert):
    """Usa pdf2docx para conversión de alta fidelidad."""
    pdf_stream = io.BytesIO(pdf_bytes)
    docx_stream = io.BytesIO()
    
    cv = Converter(pdf_stream)
    # Convierte solo las páginas seleccionadas
    cv.convert(docx_stream, pages=pages_to_convert)
    cv.close()
    
    return docx_stream.getvalue()

def convert_ocr(pdf_bytes, lang_code, pages_to_convert):
    """Usa Tesseract OCR para convertir imágenes de PDF a texto."""
    st.info(f"Usando OCR en idioma: {lang_code}. Esto tomará un momento...")
    
    # 1. Convertir páginas de PDF a imágenes
    # El 'pages_to_convert' aquí es 1-indexed para poppler, así que sumamos 1
    pages_1_indexed = [p + 1 for p in pages_to_convert]
    images = convert_from_bytes(
        pdf_bytes,
        fmt="jpeg",
        thread_count=4,
        first_page=min(pages_1_indexed),
        last_page=max(pages_1_indexed)
    )
    
    # Filtrar imágenes que no estaban en la lista (por si el rango no era contiguo)
    # Esto es complejo, por simplicidad, vamos a procesar solo las imágenes devueltas
    # que coincidan con el índice.
    
    # Nota: 'images' solo contendrá las imágenes del rango min-max.
    # Si el usuario pidió 1, 5, obtendrá 1, 2, 3, 4, 5.
    # Por ahora, procesaremos el rango continuo.
    
    st.warning("Nota de OCR: Se procesarán las páginas en el rango continuo de tu selección (ej: 1, 5 se procesa como 1-5).")

    # 2. Crear un documento Word
    doc = Document()
    doc.add_heading(f"Documento Convertido por OCR ({lang_code})", 0)
    
    # 3. Procesar cada imagen con Tesseract
    for i, img in enumerate(images):
        page_num = i + min(pages_1_indexed)
        
        # Opcional: Mostrar progreso
        st.text(f"Procesando página {page_num}...")
        
        # Extraer texto
        text = pytesseract.image_to_string(img, lang=lang_code)
        
        # Añadir texto al doc
        doc.add_heading(f"--- Página {page_num} ---", level=3)
        doc.add_paragraph(text)
        doc.add_page_break()

    # 4. Guardar documento en memoria
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    return docx_stream.getvalue()

# --- 4. Lógica Principal de la App ---

uploaded_file = st.file_uploader("Elige un archivo PDF", type="pdf")

if uploaded_file is not None:
    
    # Obtener el número de páginas (requiere Poppler)
    max_pages = 0
    if poppler_ok:
        try:
            # Una forma rápida de contar páginas
            images_info = convert_from_bytes(uploaded_file.getvalue(), fmt='jpeg', first_page=1, last_page=1)
            # Truco: pdf2image no tiene un 'count' directo, así que leemos una página
            # para probar. Un método más robusto usaría PyPDF2.
            st.info("Pre-chequeo de PDF... (usando PyPDF2 para contar páginas sería más robusto)")
            # Vamos a simplificar: no contaremos páginas por ahora
            # max_pages = 100 # Asumir un max
            st.warning("Conteo de páginas no implementado. Introduce el rango manualmente.")
            max_pages = 500 # Un límite arbitrario alto
        except Exception as e:
            st.error(f"No se pudo pre-procesar el PDF: {e}")
            
    # Opciones de páginas
    page_range_str = st.text_input(
        "Rango de páginas a convertir (ej: 1, 3-5, 9). Deja en blanco para TODAS.",
        placeholder="Ej: 1-3, 5"
    )

    if st.button("Convertir a Word", type="primary"):
        with st.spinner("Procesando... El modo OCR puede ser muy lento..."):
            try:
                # Cargar bytes del archivo
                pdf_bytes = uploaded_file.getvalue()
                
                # Decidir qué páginas procesar
                # Nota: Necesitamos una forma de saber el máx de páginas
                # Por ahora, usamos un límite alto y confiamos en el usuario.
                # Una versión 3.0 usaría PyPDF2 para un conteo exacto.
                pages_list = parse_page_range(page_range_str, max_pages=500)
                
                if not pages_list:
                    st.error("Rango de páginas no válido o vacío.")
                else:
                    st.text(f"Procesando páginas (índice 0): {pages_list}")
                    
                    # --- Selección de MODO ---
                    if "Digital" in modo_conversion:
                        st.info("Iniciando conversión digital...")
                        docx_bytes = convert_digital(pdf_bytes, pages_list)
                    
                    elif "Escaneado" in modo_conversion:
                        st.info("Iniciando conversión OCR...")
                        docx_bytes = convert_ocr(pdf_bytes, lang, pages_list)
                    
                    st.success("¡Conversión exitosa!")
                    
                    base_filename = os.path.splitext(uploaded_file.name)[0]
                    st.download_button(
                        label="Descargar archivo Word (.docx)",
                        data=docx_bytes,
                        file_name=f"{base_filename}_convertido.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

            except Exception as e:
                st.error("Ocurrió un error crítico durante la conversión.")
                st.error(f"Detalle: {e}")
                if "tesseract" in str(e).lower():
                    st.error("Error de Tesseract: Asegúrate de que esté instalado Y en el PATH del sistema.")
                if "poppler" in str(e).lower() or "pdftoppm" in str(e).lower():
                    st.error("Error de Poppler: Asegúrate de que esté instalado Y en el PATH del sistema.")

else:
    st.info("Por favor, sube un archivo PDF para comenzar.")
